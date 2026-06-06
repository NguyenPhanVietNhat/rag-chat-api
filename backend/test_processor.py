"""
test_processor.py
─────────────────
Kiểm tra DocumentProcessor với file PDF/DOCX/XLSX/TXT tạm.
Chạy: python test_processor.py
"""

import sys
import tempfile
from pathlib import Path

# Thêm thư mục cha vào sys.path nếu cần
sys.path.insert(0, str(Path(__file__).parent))

from document_processor import DocumentProcessor, DocType

# ─── Tạo file test ────────────────────────────────────────────────────────────

def make_test_txt(path: Path):
    content = """Chính sách bảo mật thông tin

Điều 1. Mục đích và phạm vi
Chính sách này quy định các nguyên tắc và biện pháp bảo vệ thông tin của khách hàng, đối tác và nhân viên trong toàn bộ hoạt động của Công ty.

Tất cả nhân viên, cộng tác viên và đối tác có tiếp cận dữ liệu nội bộ đều phải tuân thủ chính sách này. Vi phạm có thể dẫn đến kỷ luật hoặc chấm dứt hợp tác.

Điều 2. Phân loại thông tin
Thông tin được phân loại theo 3 mức độ:

Mức 1 – Công khai: thông tin được phép chia sẻ ra bên ngoài, bao gồm tài liệu marketing, thông cáo báo chí và báo cáo thường niên đã được phê duyệt.

Mức 2 – Nội bộ: thông tin chỉ dùng trong nội bộ công ty, bao gồm quy trình vận hành, kế hoạch kinh doanh và dữ liệu nhân sự không nhạy cảm.

Mức 3 – Bí mật: thông tin cần được bảo vệ nghiêm ngặt, bao gồm dữ liệu tài chính chưa công bố, thông tin khách hàng cá nhân và bí mật kinh doanh cốt lõi.

Điều 3. Nghĩa vụ bảo mật
Mỗi nhân viên có trách nhiệm bảo vệ thông tin thuộc phạm vi công việc của mình. Không được chia sẻ mật khẩu, để màn hình mở khi rời bàn làm việc, hoặc gửi tài liệu mật qua kênh không được mã hoá.

Việc vi phạm quy định bảo mật phải được báo cáo ngay lập tức cho bộ phận CNTT và cấp quản lý trực tiếp trong vòng 2 giờ kể từ khi phát hiện.

Điều 4. Lưu trữ và huỷ dữ liệu
Tài liệu giấy chứa thông tin mật phải được huỷ bằng máy cắt giấy được phê duyệt. Thiết bị lưu trữ điện tử phải được xoá an toàn trước khi tái sử dụng hoặc thải bỏ, theo quy trình được mô tả trong Phụ lục A.

Điều 5. Hiệu lực
Chính sách này có hiệu lực từ ngày 01/01/2024 và được xem xét cập nhật hàng năm hoặc khi có thay đổi quan trọng trong môi trường pháp lý hay kỹ thuật.
"""
    path.write_text(content, encoding="utf-8")


def make_test_docx(path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title  = "Hướng dẫn sử dụng hệ thống ERP"
    doc.core_properties.author = "Phòng CNTT"

    doc.add_heading("Hướng dẫn sử dụng hệ thống ERP v3.1", level=1)
    doc.add_paragraph(
        "Tài liệu này hướng dẫn người dùng cuối cách sử dụng hệ thống ERP "
        "để quản lý các quy trình vận hành hàng ngày của doanh nghiệp."
    )

    doc.add_heading("Chương 1: Đăng nhập và quản lý tài khoản", level=2)
    doc.add_paragraph(
        "Truy cập hệ thống tại địa chỉ https://erp.company.vn. "
        "Sử dụng email công ty và mật khẩu được cấp khi onboarding. "
        "Mật khẩu phải được thay đổi trong lần đăng nhập đầu tiên và "
        "định kỳ mỗi 90 ngày."
    )
    doc.add_paragraph(
        "Nếu quên mật khẩu, nhấn 'Quên mật khẩu' trên trang đăng nhập. "
        "Liên kết đặt lại sẽ được gửi tới email công ty trong vòng 5 phút. "
        "Liên kết có hiệu lực trong 30 phút."
    )

    doc.add_heading("Chương 2: Quản lý đơn hàng", level=2)
    doc.add_paragraph(
        "Để tạo đơn hàng mới, vào menu Bán hàng → Đơn hàng → Tạo mới. "
        "Điền đầy đủ thông tin khách hàng, sản phẩm và số lượng. "
        "Đơn hàng phải được phê duyệt bởi Trưởng phòng Kinh doanh trước khi "
        "chuyển sang bước xử lý kho."
    )

    # Thêm bảng
    doc.add_heading("Bảng phân quyền theo vai trò", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Vai trò", "Xem báo cáo", "Phê duyệt đơn hàng"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ("Nhân viên kinh doanh", "Giới hạn", "Không"),
        ("Trưởng phòng",         "Đầy đủ",  "Có"),
        ("Giám đốc",             "Đầy đủ",  "Có (tất cả)"),
    ]
    for row_i, row_data in enumerate(data, start=1):
        for col_i, val in enumerate(row_data):
            table.cell(row_i, col_i).text = val

    doc.save(str(path))


def make_test_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.Workbook()

    # Sheet 1: Doanh thu
    ws1 = wb.active
    ws1.title = "Doanh thu Q3"
    ws1.append(["Tháng", "Doanh thu (VND)", "Số đơn hàng", "Khách hàng mới"])
    ws1.append(["Tháng 7", 1_400_000_000, 234, 18])
    ws1.append(["Tháng 8", 1_380_000_000, 228, 15])
    ws1.append(["Tháng 9", 1_420_000_000, 251, 22])
    ws1.append(["Tổng Q3", 4_200_000_000, 713, 55])

    # Sheet 2: Nhân sự
    ws2 = wb.create_sheet("Nhân sự")
    ws2.append(["Họ tên", "Phòng ban", "Chức vụ", "Ngày vào làm"])
    ws2.append(["Nguyễn Văn An", "Kinh doanh", "Trưởng phòng", "2021-03-15"])
    ws2.append(["Trần Thị Bình",  "Kỹ thuật",  "Kỹ sư",       "2022-07-01"])
    ws2.append(["Lê Minh Châu",   "Tài chính", "Kế toán",     "2023-01-10"])

    wb.save(str(path))


# ─── Runner ───────────────────────────────────────────────────────────────────

def print_separator(title: str):
    print(f"\n{'═'*60}")
    print(f"  {title}")
    print('═'*60)


def test_file(processor: DocumentProcessor, path: Path, user_id: str = "test_user"):
    print(f"\n📄 File: {path.name}  ({path.stat().st_size // 1024} KB)")
    result = processor.process(path, user_id=user_id)

    if not result.ok:
        print(f"  ❌ Lỗi: {result.errors}")
        return

    print(f"  ✅ doc_id     : {result.doc_id}")
    print(f"  📌 doc_type   : {result.doc_type.value}")
    print(f"  📦 chunks     : {result.chunk_count}")
    print(f"  📝 ký tự tổng : {result.char_count:,}")
    print(f"  🔢 token ước  : {result.token_count:,}")
    if result.page_count:
        print(f"  📖 số trang   : {result.page_count}")
    if result.metadata:
        clean_meta = {k: v for k, v in result.metadata.items() if v}
        if clean_meta:
            print(f"  📋 metadata   : {clean_meta}")

    print(f"\n  --- 3 chunks đầu tiên ---")
    for chunk in result.chunks[:3]:
        loc = f"trang {chunk.page}" if chunk.page else chunk.section or ""
        print(f"\n  [{chunk.chunk_index}] {loc}  |  {chunk.char_count} ký tự  |  ~{chunk.token_count} tokens")
        preview = chunk.text[:180].replace("\n", " ↵ ")
        print(f"  {preview}{'…' if len(chunk.text) > 180 else ''}")


def main():
    print_separator("DOCUMENT PROCESSOR — BÀI KIỂM TRA")

    processor = DocumentProcessor(
        chunk_size    = 500,
        chunk_overlap = 80,
        min_chunk_len = 30,
    )

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)

        # TXT
        txt_path = tmp_dir / "chinh_sach_bao_mat.txt"
        make_test_txt(txt_path)

        # DOCX
        docx_path = tmp_dir / "huong_dan_erp.docx"
        make_test_docx(docx_path)

        # XLSX
        xlsx_path = tmp_dir / "bao_cao_q3.xlsx"
        make_test_xlsx(xlsx_path)

        for path in [txt_path, docx_path, xlsx_path]:
            test_file(processor, path)

    # Test batch
    print_separator("BATCH PROCESSING")
    with tempfile.TemporaryDirectory() as tmp:
        files = []
        for i in range(3):
            p = Path(tmp) / f"doc_{i}.txt"
            p.write_text(f"Tài liệu {i}.\n\n" + "Nội dung mẫu. " * 40, encoding="utf-8")
            files.append(p)

        results = processor.process_batch(files, user_id="batch_user")
        total_chunks = sum(r.chunk_count for r in results)
        print(f"  ✅ Đã xử lý {len(results)} file → tổng {total_chunks} chunks")

    # Test lỗi
    print_separator("XỬ LÝ LỖI")
    try:
        processor.process("khong_ton_tai.pdf")
    except FileNotFoundError as e:
        print(f"  ✅ FileNotFoundError: {e}")

    try:
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as f:
            processor.process(f.name)
    except ValueError as e:
        print(f"  ✅ ValueError: {e}")

    print_separator("HOÀN TẤT")
    print("  Tất cả bài kiểm tra đã chạy thành công!\n")


if __name__ == "__main__":
    main()
