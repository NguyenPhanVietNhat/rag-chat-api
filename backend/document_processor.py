"""
document_processor.py
─────────────────────
Đọc tài liệu (PDF, DOCX, XLSX, TXT) và chia thành chunks có metadata.
Dùng trong pipeline RAG trước bước embedding.

Cách dùng:
    processor = DocumentProcessor()
    result    = processor.process("bao_cao.pdf", user_id="u_001")
    chunks    = result.chunks   # List[Chunk]
"""

from __future__ import annotations

import hashlib
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional

# ── PDF ──────────────────────────────────────────────────────────────────────
import pdfplumber                          # text + table extraction
from pypdf import PdfReader                # metadata + fallback

# ── DOCX ─────────────────────────────────────────────────────────────────────
from docx import Document as DocxDocument  # python-docx

# ── XLSX ─────────────────────────────────────────────────────────────────────
import openpyxl

# ── Chunking ─────────────────────────────────────────────────────────────────
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ═══════════════════════════════════════════════════════════════════════════════
# DATA MODELS
# ═══════════════════════════════════════════════════════════════════════════════

class DocType(str, Enum):
    PDF  = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    TXT  = "txt"
    UNKNOWN = "unknown"


@dataclass
class Chunk:
    """Một đoạn văn bản đã được xử lý, sẵn sàng để embed."""

    chunk_id:    str          # hash duy nhất cho chunk này
    doc_id:      str          # hash của file gốc
    text:        str          # nội dung chunk (đã clean)
    token_count: int          # ước lượng số token
    char_count:  int

    # Metadata định vị nguồn gốc
    page:        Optional[int]  = None   # trang trong PDF / sheet index
    section:     Optional[str]  = None   # heading / sheet name / row range
    chunk_index: int            = 0      # thứ tự trong tài liệu

    # Metadata tài liệu
    source_file: str            = ""
    doc_type:    str            = ""
    user_id:     str            = ""
    created_at:  str            = ""


@dataclass
class ProcessResult:
    """Kết quả trả về sau khi xử lý một tài liệu."""

    doc_id:      str
    source_file: str
    doc_type:    DocType
    chunks:      list[Chunk]       = field(default_factory=list)
    page_count:  int               = 0
    char_count:  int               = 0
    token_count: int               = 0
    errors:      list[str]         = field(default_factory=list)
    metadata:    dict              = field(default_factory=dict)

    @property
    def chunk_count(self) -> int:
        return len(self.chunks)

    @property
    def ok(self) -> bool:
        return len(self.chunks) > 0


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT CLEANER
# ═══════════════════════════════════════════════════════════════════════════════

class TextCleaner:
    """Chuẩn hoá văn bản tiếng Việt / tiếng Anh trước khi chunk."""

    # Ký tự đặc biệt hay xuất hiện trong PDF xuất kém
    _LIGATURE_MAP = str.maketrans({
        "\ufb01": "fi", "\ufb02": "fl", "\ufb00": "ff",
        "\ufb03": "ffi", "\ufb04": "ffl",
        "\u2013": "-", "\u2014": "-",   # en-dash, em-dash
        "\u2018": "'", "\u2019": "'",   # curly quotes
        "\u201c": '"', "\u201d": '"',
        "\u00a0": " ",                  # non-breaking space
        "\t":     " ",
    })

    def clean(self, text: str) -> str:
        if not text:
            return ""

        # 1. Unicode NFC + thay ligature
        text = unicodedata.normalize("NFC", text)
        text = text.translate(self._LIGATURE_MAP)

        # 2. Xoá ký tự điều khiển (trừ newline)
        text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)

        # 3. Gộp nhiều dòng trống → tối đa 2 newline
        text = re.sub(r"\n{3,}", "\n\n", text)

        # 4. Gộp dấu cách thừa trong dòng
        text = re.sub(r"[ ]{2,}", " ", text)

        # 5. Xoá dòng chỉ có số trang / header/footer lặp lại
        text = re.sub(r"(?m)^\s*-?\s*\d+\s*-?\s*$", "", text)

        # 6. Strip từng dòng
        lines = [ln.rstrip() for ln in text.splitlines()]
        text = "\n".join(lines).strip()

        return text

    def estimate_tokens(self, text: str) -> int:
        """Ước tính token theo heuristic 1 token ≈ 4 ký tự (tiếng Anh)
        hoặc ≈ 2 ký tự (tiếng Việt / CJK). Dùng trung bình."""
        latin  = sum(1 for c in text if ord(c) < 0x300)
        others = len(text) - latin
        return max(1, latin // 4 + others // 2)


# ═══════════════════════════════════════════════════════════════════════════════
# READERS  (một class cho mỗi định dạng)
# ═══════════════════════════════════════════════════════════════════════════════

class PdfReader_:
    """Đọc PDF bằng pdfplumber (ưu tiên) + pypdf (fallback + metadata)."""

    def __init__(self, cleaner: TextCleaner):
        self.cleaner = cleaner

    def read(self, path: Path) -> tuple[list[dict], dict]:
        """
        Trả về:
            pages    – list[{"page": int, "text": str}]
            metadata – dict
        """
        pages: list[dict] = []
        metadata: dict    = {}

        # ── pypdf: metadata ───────────────────────────────────────────────
        try:
            reader   = PdfReader(str(path))
            info     = reader.metadata or {}
            metadata = {
                "title":    info.get("/Title",    ""),
                "author":   info.get("/Author",   ""),
                "subject":  info.get("/Subject",  ""),
                "creator":  info.get("/Creator",  ""),
                "pages":    len(reader.pages),
            }
        except Exception:
            pass

        # ── pdfplumber: text + tables ──────────────────────────────────────
        try:
            with pdfplumber.open(str(path)) as pdf:
                metadata["pages"] = metadata.get("pages") or len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    parts: list[str] = []

                    # 1. Extract text thường
                    raw = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                    if raw.strip():
                        parts.append(raw)

                    # 2. Extract bảng → chuyển thành text dạng CSV-like
                    for table in page.extract_tables():
                        rows = []
                        for row in table:
                            cells = [str(c or "").strip() for c in row]
                            if any(cells):
                                rows.append(" | ".join(cells))
                        if rows:
                            parts.append("\n".join(rows))

                    combined = self.cleaner.clean("\n".join(parts))
                    if combined:
                        pages.append({"page": page_num, "text": combined})

        except Exception as e:
            # Fallback: dùng pypdf thuần để lấy text
            try:
                reader = PdfReader(str(path))
                for i, pg in enumerate(reader.pages, start=1):
                    raw = pg.extract_text() or ""
                    cleaned = self.cleaner.clean(raw)
                    if cleaned:
                        pages.append({"page": i, "text": cleaned})
            except Exception as e2:
                raise RuntimeError(f"Không đọc được PDF: {e2}") from e2

        return pages, metadata


class DocxReader:
    """Đọc DOCX: giữ cấu trúc heading / paragraph / table."""

    def __init__(self, cleaner: TextCleaner):
        self.cleaner = cleaner

    def read(self, path: Path) -> tuple[list[dict], dict]:
        doc      = DocxDocument(str(path))
        sections: list[dict] = []
        metadata: dict = {
            "title":  doc.core_properties.title  or "",
            "author": doc.core_properties.author or "",
        }

        current_heading = "Phần đầu tài liệu"
        buffer: list[str] = []

        def flush():
            nonlocal buffer
            text = self.cleaner.clean("\n".join(buffer))
            if text:
                sections.append({"section": current_heading, "text": text})
            buffer = []

        for para in doc.paragraphs:
            style = para.style.name or ""
            text  = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                flush()
                current_heading = text
            else:
                buffer.append(text)

        # Tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                buffer.append("\n".join(rows))

        flush()
        return sections, metadata


class XlsxReader:
    """Đọc XLSX: mỗi sheet → một section, mỗi hàng → một dòng text."""

    def __init__(self, cleaner: TextCleaner):
        self.cleaner = cleaner

    def read(self, path: Path) -> tuple[list[dict], dict]:
        wb       = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections: list[dict] = []
        metadata = {"sheets": wb.sheetnames}

        for sheet_name in wb.sheetnames:
            ws   = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                # Bỏ hàng rỗng hoàn toàn
                if any(c for c in cells):
                    rows.append(" | ".join(cells))

            if rows:
                # Gộp mỗi BATCH_ROWS hàng thành một section để tránh chunk quá lớn
                BATCH = 80
                for i in range(0, len(rows), BATCH):
                    batch_text = self.cleaner.clean("\n".join(rows[i:i+BATCH]))
                    if batch_text:
                        label = f"{sheet_name} (hàng {i+1}–{min(i+BATCH, len(rows))})"
                        sections.append({"section": label, "text": batch_text})

        wb.close()
        return sections, metadata


class TxtReader:
    def __init__(self, cleaner: TextCleaner):
        self.cleaner = cleaner

    def read(self, path: Path) -> tuple[list[dict], dict]:
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                raw = path.read_text(encoding=enc)
                return [{"text": self.cleaner.clean(raw)}], {}
            except UnicodeDecodeError:
                continue
        raise RuntimeError(f"Không đọc được file text: {path.name}")


# ═══════════════════════════════════════════════════════════════════════════════
# CHUNKER
# ═══════════════════════════════════════════════════════════════════════════════

class Chunker:
    """
    Chia text thành chunks bằng RecursiveCharacterTextSplitter.
    Dùng cấu trúc ưu tiên: đoạn văn → câu → từ → ký tự.
    """

    # Separators theo thứ tự ưu tiên (tiếng Việt + tiếng Anh)
    _SEPARATORS = [
        "\n\n",      # Đoạn văn
        "\n",        # Xuống dòng
        "。", "！", "？",   # Câu tiếng Nhật/Trung (nếu có)
        ". ", "! ", "? ",  # Câu tiếng Anh
        "। ",              # Câu tiếng Hindi
        " ",               # Từ
        "",                # Ký tự
    ]

    def __init__(
        self,
        chunk_size:    int = 500,   # ký tự tối đa mỗi chunk
        chunk_overlap: int = 80,    # overlap để giữ ngữ cảnh
    ):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=self._SEPARATORS,
            is_separator_regex=False,
            keep_separator=True,
        )
        self.chunk_size    = chunk_size
        self.chunk_overlap = chunk_overlap

    def split(self, text: str) -> list[str]:
        """Chia text → list[str]. Mỗi phần tử là 1 chunk."""
        if not text or not text.strip():
            return []
        return [c for c in self.splitter.split_text(text) if c.strip()]


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT PROCESSOR  (entry point)
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentProcessor:
    """
    Bộ xử lý tài liệu chính.
    Tự nhận diện định dạng file, đọc nội dung, làm sạch và chia chunks.

    Ví dụ:
        processor = DocumentProcessor(chunk_size=500, chunk_overlap=80)
        result    = processor.process("tài_liệu.pdf", user_id="u_123")

        for chunk in result.chunks:
            print(chunk.chunk_index, chunk.page, chunk.text[:80])
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf":  DocType.PDF,
        ".docx": DocType.DOCX,
        ".doc":  DocType.DOCX,
        ".xlsx": DocType.XLSX,
        ".xls":  DocType.XLSX,
        ".txt":  DocType.TXT,
        ".md":   DocType.TXT,
        ".csv":  DocType.TXT,
    }

    def __init__(
        self,
        chunk_size:    int = 500,
        chunk_overlap: int = 80,
        min_chunk_len: int = 30,   # bỏ chunk quá ngắn (header, số trang...)
    ):
        self.cleaner       = TextCleaner()
        self.chunker       = Chunker(chunk_size, chunk_overlap)
        self.min_chunk_len = min_chunk_len

        # Lazy-init readers
        self._readers: dict[DocType, object] = {}

    # ── Private helpers ───────────────────────────────────────────────────────

    def _get_reader(self, doc_type: DocType):
        if doc_type not in self._readers:
            cls_map = {
                DocType.PDF:  PdfReader_,
                DocType.DOCX: DocxReader,
                DocType.XLSX: XlsxReader,
                DocType.TXT:  TxtReader,
            }
            self._readers[doc_type] = cls_map[doc_type](self.cleaner)
        return self._readers[doc_type]

    def _make_doc_id(self, path: Path) -> str:
        """SHA-256 của tên file + kích thước + mtime."""
        stat = path.stat()
        raw  = f"{path.name}:{stat.st_size}:{stat.st_mtime}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]

    def _make_chunk_id(self, doc_id: str, index: int, text: str) -> str:
        raw = f"{doc_id}:{index}:{text[:64]}"
        return hashlib.sha256(raw.encode()).hexdigest()[:12]

    def _sections_to_chunks(
        self,
        sections: list[dict],
        doc_id: str,
        source_file: str,
        doc_type: DocType,
        user_id: str,
    ) -> list[Chunk]:
        """Nhận sections từ reader → split → tạo Chunk objects."""
        chunks:  list[Chunk] = []
        counter: int         = 0
        now = datetime.utcnow().isoformat()

        for sec in sections:
            text    = sec.get("text", "")
            page    = sec.get("page")
            section = sec.get("section")

            parts = self.chunker.split(text)
            for part in parts:
                if len(part) < self.min_chunk_len:
                    continue

                tokens = self.cleaner.estimate_tokens(part)
                chunk  = Chunk(
                    chunk_id    = self._make_chunk_id(doc_id, counter, part),
                    doc_id      = doc_id,
                    text        = part,
                    token_count = tokens,
                    char_count  = len(part),
                    page        = page,
                    section     = section,
                    chunk_index = counter,
                    source_file = source_file,
                    doc_type    = doc_type.value,
                    user_id     = user_id,
                    created_at  = now,
                )
                chunks.append(chunk)
                counter += 1

        return chunks

    # ── Public API ────────────────────────────────────────────────────────────

    def detect_type(self, path: Path) -> DocType:
        ext = path.suffix.lower()
        return self.SUPPORTED_EXTENSIONS.get(ext, DocType.UNKNOWN)

    def process(
        self,
        file_path: str | Path,
        user_id:   str = "anonymous",
    ) -> ProcessResult:
        """
        Xử lý một file tài liệu.

        Tham số:
            file_path – đường dẫn tới file (str hoặc Path)
            user_id   – ID người dùng (dùng để lọc trong Vector DB)

        Trả về ProcessResult với:
            .ok         – True nếu có ít nhất 1 chunk
            .chunks     – List[Chunk]
            .chunk_count
            .errors     – List[str] nếu có lỗi không nghiêm trọng
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {path}")

        doc_type = self.detect_type(path)
        if doc_type == DocType.UNKNOWN:
            raise ValueError(
                f"Định dạng '{path.suffix}' không được hỗ trợ. "
                f"Hỗ trợ: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        doc_id = self._make_doc_id(path)
        result = ProcessResult(
            doc_id      = doc_id,
            source_file = path.name,
            doc_type    = doc_type,
        )

        try:
            reader   = self._get_reader(doc_type)
            sections, metadata = reader.read(path)
            result.metadata = metadata

            # PDF: sections là list[{page, text}]
            # Đặt page_count
            if doc_type == DocType.PDF:
                result.page_count = metadata.get("pages", len(sections))

            result.chunks = self._sections_to_chunks(
                sections, doc_id, path.name, doc_type, user_id
            )
            result.char_count  = sum(c.char_count  for c in result.chunks)
            result.token_count = sum(c.token_count for c in result.chunks)

        except Exception as exc:
            result.errors.append(str(exc))

        return result

    def process_batch(
        self,
        file_paths: list[str | Path],
        user_id:    str = "anonymous",
    ) -> list[ProcessResult]:
        """Xử lý nhiều file cùng lúc. Lỗi từng file không dừng toàn bộ batch."""
        results = []
        for fp in file_paths:
            try:
                results.append(self.process(fp, user_id))
            except Exception as e:
                results.append(ProcessResult(
                    doc_id      = "",
                    source_file = str(fp),
                    doc_type    = DocType.UNKNOWN,
                    errors      = [str(e)],
                ))
        return results
